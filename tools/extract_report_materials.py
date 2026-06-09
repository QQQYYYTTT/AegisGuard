from pathlib import Path
import json
import re

from docx import Document
from pypdf import PdfReader


ROOT = Path(r"F:\2026信安赛")
OUT = Path(r"F:\2026信安赛\AegisGuard\build\report_materials")

FILES = {
    "current_report": ROOT / r"信安赛初赛\作品报告_v9.docx",
    "langgraph": ROOT / r"信安赛初赛\测试\langraph-测试结果.md",
    "dpi_cost": ROOT / r"信安赛初赛\测试\Token消耗+时延（DPI攻击）.md",
    "generalization": ROOT / r"信安赛初赛\测试\泛化能力测试（不同底层模型）.md",
    "gaiguardian": ROOT / r"导师\往年作品\PDF\GAIGuardian：面向生成式AI的动态演化监管系统（2025年全国一等奖）（内部学习资料，勿上网外传）.pdf",
}


def clean(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def docx_to_text(path: Path):
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for i, table in enumerate(doc.tables, 1):
        parts.append(f"\n[表格 {i}]")
        for row in table.rows:
            cells = [clean(c.text) for c in row.cells]
            parts.append(" | ".join(cells))
    return clean("\n".join(parts))


def pdf_to_text(path: Path, max_pages=None):
    reader = PdfReader(str(path))
    pages = []
    for idx, page in enumerate(reader.pages):
        if max_pages is not None and idx >= max_pages:
            break
        text = page.extract_text() or ""
        pages.append(f"\n--- page {idx + 1} ---\n{text}")
    return clean("\n".join(pages))


def read_text(path: Path):
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def headings(text: str):
    hs = []
    for line in text.splitlines():
        s = line.strip()
        if re.match(r"^(第[一二三四五六七八九十]+章|[一二三四五六七八九十]+、|\d+(\.\d+){0,3}\s+|#+\s+)", s):
            hs.append(s[:160])
    return hs


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    summary = {}

    current = docx_to_text(FILES["current_report"])
    (OUT / "current_report.txt").write_text(current, encoding="utf-8")
    summary["current_report_headings"] = headings(current)
    summary["current_report_chars"] = len(current)

    for key in ("langgraph", "dpi_cost", "generalization"):
        text = read_text(FILES[key])
        (OUT / f"{key}.txt").write_text(text, encoding="utf-8")
        summary[f"{key}_headings"] = headings(text)
        summary[f"{key}_chars"] = len(text)

    gai = pdf_to_text(FILES["gaiguardian"])
    (OUT / "gaiguardian.txt").write_text(gai, encoding="utf-8")
    summary["gaiguardian_headings"] = headings(gai)
    summary["gaiguardian_chars"] = len(gai)

    (OUT / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
