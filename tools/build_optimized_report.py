from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path(r"F:\2026信安赛\信安赛初赛\作品报告_v9.docx")
OUT_DIR = Path(r"F:\2026信安赛\AegisGuard\build\optimized_report")
OUT_DOCX = OUT_DIR / "AegisGuard_作品报告_v10_新增实验整合版.docx"
OUT_PLAN = OUT_DIR / "新增实验整合方案.md"


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p._element = new_p
    if style is not None:
        p.style = style
    if text:
        p.add_run(text)
    return p


def insert_table_after(paragraph, rows):
    tbl = paragraph._parent.add_table(rows=0, cols=len(rows[0]), width=Inches(6.5))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for r, row in enumerate(rows):
        cells = tbl.add_row().cells
        for c, value in enumerate(row):
            cells[c].text = str(value)
            cells[c].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[c].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c != 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    if r == 0:
                        run.bold = True
    paragraph._p.addnext(tbl._tbl)
    spacer = OxmlElement("w:p")
    tbl._tbl.addnext(spacer)
    p = paragraph._parent.add_paragraph()
    p._p = spacer
    p._element = spacer
    return p


def paragraph_with_style(doc, text, style_name):
    p = doc.add_paragraph(text)
    try:
        p.style = style_name
    except KeyError:
        pass
    return p


def append_after(anchor, items):
    cur = anchor
    for item in items:
        kind = item[0]
        if kind == "p":
            cur = insert_paragraph_after(cur, item[1], item[2] if len(item) > 2 else None)
            for run in cur.runs:
                run.font.size = Pt(10.5)
        elif kind == "h":
            cur = insert_paragraph_after(cur, item[1], item[2])
        elif kind == "caption":
            cur = insert_paragraph_after(cur, item[1], None)
            cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cur.runs:
                run.font.size = Pt(9)
                run.bold = True
        elif kind == "table":
            cur = insert_table_after(cur, item[1])
    return cur


def set_update_fields(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def find_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"paragraph not found: {text}")


def add_intro_plan():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    plan = """# AegisGuard 作品报告新增实验整合方案

## 一、总体判断

现有报告的测试章节已经具备“测试目标、测试方案、测试环境、功能安全测试、性能开销、结果总结”的基本框架，不建议推翻重写。新增数据应作为第三章的增强证据，重点补齐三个原报告相对薄弱的点：

1. 真实 Agent 框架下的闭环验证：用 LangGraph Financial Agent 的 200 case 结果证明系统并非只在模拟流程中有效。
2. 跨底层模型泛化：用 5 个模型、750 case 的 DPI 测试证明 AegisGuard 的防护效果不依赖单一模型。
3. 工程开销：用 Token 与时延数据说明系统为了运行时监管引入了可度量开销，并给出优化方向。

## 二、推荐放置位置

- 摘要：加入“200 个真实 Agent case ASR 从 7.50% 降至 0.00%、任务成功率从 85.50% 提升至 92.50%”等量化结果。
- 3.4 功能测试与安全效果分析：保留原有单点攻击场景，用作机制验证。
- 新增 3.6 真实 Agent 安全有效性测试：放 LangGraph 六类场景总表，作为第三章主实证结果。
- 新增 3.7 跨模型泛化能力测试：放 5 个底层模型对比表，强调与模型无关的运行时门控能力。
- 新增 3.8 Token 消耗与时延开销测试：放 15 case 小样本开销表，作为工程成本分析，不作为主安全结论。
- 原 3.6 测试结果总结顺延为 3.9，并吸收新增实验结论。
- 第五章：补充当前阶段成果，同时诚实说明 OPI 保守拒绝、MP memory_found 未命中、部分模型补跑合并等局限。

## 三、测试逻辑是否需要调整

需要小幅调整，不需要重写。建议将第三章测试逻辑表述为四层递进：

1. 机制正确性测试：验证 Message Gate、Action Gate、Return Gate、授权链路、记忆沙箱是否按设计工作。
2. 攻击有效性对照测试：在无防御组与 AegisGuard 组之间比较 ASR、任务完成率、拒绝率。
3. 泛化测试：更换底层模型，保持 Agent、任务、攻击族和评价指标不变。
4. 工程代价测试：用 Token、耗时、拒绝率评价可部署性。

这样写会更接近往年获奖作品常见的“需求痛点-关键技术-分层实验-结果分析-局限展望”结构，语言上也更稳健。

## 四、写作注意事项

- 不要把 OPI 和 Mixed 的 0% ASR 写成 AegisGuard 额外降低攻击成功率，因为无防御组本身未被打穿；应写成“保持 0% ASR，并观察可用性影响”。
- MP 结果中 memory_found 均为 0，应避免声称已经完整验证“记忆检索命中后的污染防护”；更严谨的写法是“当前 MP/combined_attack 配置下消除了观测到的成功攻击样本，后续将补充更强命中样本”。
- Token/时延小样本只有 15 case，适合说明开销，不适合替代 75 case/750 case 主表。
- Doubao 和 GLM 的 AegisGuard 结果包含补跑合并，应在表下注明，增强可信度。
"""
    OUT_PLAN.write_text(plan, encoding="utf-8")


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(SRC))

    abstract_anchor = doc.paragraphs[118]
    abstract_addition = (
        "在最新真实 Agent 实验中，本作品进一步基于 LangGraph Financial Agent 构建了 DPI、OPI、Mixed、MP、POT 与 Clean 六类场景的对照测试。"
        "实验共覆盖 200 个真实 Agent case，AegisGuard 将总体攻击成功率由 7.50% 降至 0.00%，在攻击场景中将任务成功率由 88.00% 提升至 93.14%，在包含无攻击 Clean 场景后总体任务成功率由 85.50% 提升至 92.50%。"
        "同时，面向 deepseek-v4-flash、qwen3.5-plus、doubao-seed-1.6、glm-4.7-flash 与 gpt-4o-mini 五类底层模型的 750 case 泛化测试显示，AegisGuard 均能将 DPI 攻击成功率压降至 0.00%，并在各模型上保持或提升任务完成率，说明系统具备较好的框架适配性、模型泛化性与工程落地价值。"
    )
    insert_paragraph_after(abstract_anchor, abstract_addition)

    # Rename old summary section after inserting new evidence sections.
    renames = {
        "3.6测试结果总结": "3.9 测试结果总结",
        "3.6.1 核心结果汇总": "3.9.1 核心结果汇总",
        "3.6.2有效性分析": "3.9.2 有效性分析",
        "3.6.3局限性分析": "3.9.3 局限性分析",
        "3.6测试结果总结\t76": "3.6 真实 Agent 安全有效性测试\t76",
        "3.6.1 核心结果汇总\t76": "3.7 跨底层模型泛化能力测试\t77",
        "3.6.2有效性分析\t77": "3.8 Token 消耗与时延开销测试\t78",
        "3.6.3局限性分析\t77": "3.9 测试结果总结\t79",
    }
    for p in doc.paragraphs:
        key = p.text.strip()
        if key in renames:
            p.text = renames[key]

    anchor_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "3.9 测试结果总结")
    anchor = doc.paragraphs[anchor_index - 1]
    sections = [
        ("h", "3.6 真实 Agent 安全有效性测试", "Heading 2"),
        ("p", "为验证 AegisGuard 在真实 Agent 运行链路中的防护效果，本节在 ASB-native LangGraph Financial Agent 上构建对照实验。测试对象覆盖直接提示注入（DPI）、观测提示注入（OPI）、DPI 与 OPI 组合攻击（Mixed）、记忆污染链路（MP）、计划思维后门攻击（POT）以及无攻击 Clean 基线。每类实验均设置无防御组与 AegisGuard 防御组，统一统计攻击成功率（ASR）、原始任务成功率、拒绝率和运行耗时。"),
        ("caption", "表 3-新增1 LangGraph Financial Agent 真实 Agent 测试结果"),
        ("table", [
            ["攻击类型", "组别", "样本数", "攻击成功数", "ASR", "原始任务成功率", "拒绝率", "运行耗时"],
            ["DPI", "无防御", 75, 13, "17.33%", "77.33%", "1.33%", "1395.856s"],
            ["DPI", "AegisGuard", 75, 0, "0.00%", "90.67%", "0.00%", "1528.343s"],
            ["OPI", "无防御", 25, 0, "0.00%", "100.00%", "0.00%", "468.663s"],
            ["OPI", "AegisGuard", 25, 0, "0.00%", "88.00%", "12.00%", "372.126s"],
            ["Mixed", "无防御", 25, 0, "0.00%", "92.00%", "0.00%", "306.302s"],
            ["Mixed", "AegisGuard", 25, 0, "0.00%", "96.00%", "0.00%", "310.128s"],
            ["MP", "无防御", 25, 1, "4.00%", "96.00%", "0.00%", "341.779s"],
            ["MP", "AegisGuard", 25, 0, "0.00%", "100.00%", "0.00%", "342.211s"],
            ["POT", "无防御", 25, 1, "4.00%", "96.00%", "0.00%", "463.076s"],
            ["POT", "AegisGuard", 25, 0, "0.00%", "96.00%", "0.00%", "312.780s"],
            ["Clean", "无防御", 25, 0, "0.00%", "68.00%", "12.00%", "392.363s"],
            ["Clean", "AegisGuard", 25, 0, "0.00%", "88.00%", "12.00%", "354.102s"],
        ]),
        ("p", "从实验结果看，AegisGuard 在全部 200 个真实 Agent case 中将攻击成功数从 15 降至 0，总体 ASR 从 7.50% 降至 0.00%。在 DPI 场景中，ASR 从 17.33% 降至 0.00%，是安全收益最显著的场景；在 MP 与 POT 场景中，无防御组各出现 1 个成功攻击样本，AegisGuard 均将其压降为 0；在 OPI 与 Mixed 场景中，无防御组本身未被当前样本打穿，AegisGuard 主要体现为继续保持 0% ASR。"),
        ("p", "可用性方面，攻击场景下任务成功率由 88.00% 提升至 93.14%，包含 Clean 场景后总体任务成功率由 85.50% 提升至 92.50%。这表明 AegisGuard 的防护策略并非通过简单拒绝实现安全，而是在阻断攻击目标的同时维持了较高任务完成能力。需要注意的是，OPI 场景中 AegisGuard 组拒绝率达到 12.00%，说明在部分外部观测内容场景下策略仍偏保守，后续可进一步通过风险阈值、工具类别和任务上下文联合调优降低误拒。"),
        ("caption", "表 3-新增2 真实 Agent 测试总体汇总"),
        ("table", [
            ["范围", "组别", "样本数", "攻击成功数", "总 ASR", "任务成功率", "拒绝率"],
            ["攻击场景（DPI/OPI/Mixed/MP/POT）", "无防御", 175, 15, "8.57%", "88.00%", "0.57%"],
            ["攻击场景（DPI/OPI/Mixed/MP/POT）", "AegisGuard", 175, 0, "0.00%", "93.14%", "1.71%"],
            ["全部场景（含 Clean）", "无防御", 200, 15, "7.50%", "85.50%", "2.00%"],
            ["全部场景（含 Clean）", "AegisGuard", 200, 0, "0.00%", "92.50%", "3.00%"],
        ]),
        ("h", "3.7 跨底层模型泛化能力测试", "Heading 2"),
        ("p", "为排除防护效果依赖单一底层模型的可能性，本节进一步固定 Agent 框架、任务集合、攻击族和评价指标，仅替换底层大模型，选取 deepseek-v4-flash、qwen3.5-plus、doubao-seed-1.6、glm-4.7-flash 与 gpt-4o-mini 五类模型进行 DPI 攻击泛化测试。每个模型在无防御与 AegisGuard 两种模式下各测试 75 个 case，总计 750 个 case。"),
        ("caption", "表 3-新增3 不同底层模型下的 DPI 泛化测试结果"),
        ("table", [
            ["模型", "无防御 ASR", "AegisGuard ASR", "ASR 降低", "无防御任务成功率", "AegisGuard 任务成功率", "Token 开销"],
            ["deepseek-v4-flash", "17.33%", "0.00%", "17.33pp", "73.33%", "90.67%", "+14.52%"],
            ["qwen3.5-plus", "17.33%", "0.00%", "17.33pp", "76.00%", "90.67%", "+14.54%"],
            ["doubao-seed-1.6", "17.33%", "0.00%", "17.33pp", "80.00%", "93.33%", "+14.75%"],
            ["glm-4.7-flash", "17.33%", "0.00%", "17.33pp", "77.33%", "88.00%", "+14.40%"],
            ["gpt-4o-mini", "14.67%", "0.00%", "14.67pp", "76.00%", "88.00%", "+14.73%"],
        ]),
        ("p", "结果显示，在五类模型上，无防御组 DPI ASR 介于 14.67% 至 17.33% 之间，均存在不同程度的攻击成功；接入 AegisGuard 后，五类模型的 ASR 均降至 0.00%。同时，任务成功率在五类模型上均出现提升，提升幅度介于 10.67 个百分点至 17.34 个百分点之间。该结果说明 AegisGuard 的核心防护能力主要来自运行时授权链路、门控决策与上下文隔离机制，而不是依赖某一模型自身的拒答倾向或安全微调能力。"),
        ("p", "实验中 doubao-seed-1.6 与 glm-4.7-flash 的 AegisGuard 组使用了补跑合并结果，其中 GLM 组时间列排除了失败 timeout chunk。因此，跨模型实验更适合作为安全效果与 Token 开销的证据，时延对比需结合网络中转、模型服务稳定性和补跑策略共同解释。"),
        ("h", "3.8 Token 消耗与时延开销测试", "Heading 2"),
        ("p", "AegisGuard 在执行链路中引入 Message Gate、Action Gate、Return Gate 与审计记录等运行时检查，理论上会增加一定上下文与处理开销。为评估该开销，本节在 LangGraph Financial Agent、gpt-4o-mini 与 DPI 场景下选取 15 个 case 进行 Token 消耗与时延对比。"),
        ("caption", "表 3-新增4 DPI 场景下 Token 消耗与时延对比"),
        ("table", [
            ["模式", "样本数", "Input Tokens", "Output Tokens", "Total Tokens", "Total Time", "平均 Tokens/Case", "平均时延/Case"],
            ["无防御", 15, "43,877", "6,342", "50,219", "994.589s", "3347.93", "66.31s"],
            ["AegisGuard", 15, "50,910", "6,448", "57,358", "1448.112s", "3823.87", "96.54s"],
            ["变化率", "-", "+16.03%", "+1.67%", "+14.22%", "+45.60%", "+14.22%", "+45.60%"],
        ]),
        ("p", "小样本结果显示，AegisGuard 相比无防御组总 Token 增加 14.22%，总耗时增加 45.60%。其中 Token 增长主要来自安全策略上下文、门控判定提示与审计信息组织；时延增长则与额外门控调用、模型服务响应波动和串行执行链路有关。由于该组仅覆盖 15 个 case，报告中将其定位为工程开销估计，不作为安全有效性的主证据。后续优化可从策略提示压缩、门控结果缓存、低风险动作快速通道和异步审计写入四个方向降低部署成本。"),
    ]
    append_after(anchor, sections)

    conclusion_anchor = find_para(doc, "5.2 当前阶段成果")
    append_after(conclusion_anchor, [
        ("p", "在最新测试阶段，系统已完成真实 Agent 框架和跨模型泛化两类补充实验：一是在 LangGraph Financial Agent 上完成 200 个 case 的安全有效性测试，验证 AegisGuard 能够在真实任务链路中将总体 ASR 从 7.50% 降至 0.00%；二是在五类底层模型上完成 750 个 case 的 DPI 泛化测试，验证系统在不同模型上的攻击压制效果均保持稳定。上述结果进一步支撑了本作品“面向 Agent 运行时，而非单一模型输出过滤”的技术定位。"),
    ])

    limitation_anchor = find_para(doc, "5.3 现阶段仍存在的不足")
    append_after(limitation_anchor, [
        ("p", "从新增实验看，系统仍存在三点需要继续完善：其一，OPI 场景中出现一定保守拒绝，说明门控策略在外部观测内容可信度评估上仍有调参空间；其二，当前 MP 测试批次中 memory_found 均为 0，尚需补充更强记忆检索命中样本来进一步验证记忆沙箱能力；其三，跨模型泛化实验中部分模型结果采用补跑合并方式，后续应在更稳定的服务环境下扩展样本规模并统一时延统计口径。"),
    ])

    set_update_fields(doc)
    doc.save(str(OUT_DOCX))


if __name__ == "__main__":
    add_intro_plan()
    build()
    print(OUT_DOCX)
    print(OUT_PLAN)
