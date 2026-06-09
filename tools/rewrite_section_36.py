from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path(r"F:\2026信安赛\AegisGuard\build\optimized_report\AegisGuard_作品报告_v10_新增实验整合版.docx")
OUT = Path(r"F:\2026信安赛\AegisGuard\build\optimized_report\AegisGuard_作品报告_v11_3.6细化版.docx")


def block_text(block):
    if block.tag.endswith("}p"):
        return "".join(block.xpath(".//w:t/text()")).strip()
    return ""


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    p._element = new_p
    if style is not None:
        p.style = style
    if text:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p


def insert_table_after(paragraph, rows, font_size=8.5):
    tbl = paragraph._parent.add_table(rows=0, cols=len(rows[0]), width=Inches(6.5))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for r, row in enumerate(rows):
        cells = tbl.add_row().cells
        for c, value in enumerate(row):
            cells[c].text = str(value)
            cells[c].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[c].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c != 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    if r == 0:
                        run.bold = True
    paragraph._p.addnext(tbl._tbl)
    spacer = OxmlElement("w:p")
    tbl._tbl.addnext(spacer)
    p = paragraph._parent.add_paragraph()
    p._p = spacer
    p._element = spacer
    return p


def caption_after(cur, text):
    p = insert_paragraph_after(cur, text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(9)
    return p


def append_items(anchor, items):
    cur = anchor
    for item in items:
        kind = item[0]
        if kind == "h":
            cur = insert_paragraph_after(cur, item[1], item[2])
        elif kind == "p":
            cur = insert_paragraph_after(cur, item[1])
        elif kind == "cap":
            cur = caption_after(cur, item[1])
        elif kind == "table":
            cur = insert_table_after(cur, item[1], item[2] if len(item) > 2 else 8.5)
    return cur


def set_update_fields(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def find_body_range(doc):
    body = doc.element.body
    blocks = list(body.iterchildren())
    start = None
    end = None
    hit36 = 0
    hit37 = 0
    for i, block in enumerate(blocks):
        text = block_text(block)
        if text.startswith("3.6 真实 Agent 安全有效性测试"):
            hit36 += 1
            if hit36 == 2:
                start = i
        if text.startswith("3.7 跨底层模型泛化能力测试"):
            hit37 += 1
            if hit37 == 2 and start is not None:
                end = i
                break
    if start is None or end is None:
        raise RuntimeError("未找到正文 3.6 到 3.7 的替换范围")
    return body, blocks, start, end


def detailed_items():
    attack_tool_rows = [
        ["攻击工具", "模拟风险", "业务含义", "AegisGuard 主要防护点"],
        ["InvestmentDiversion", "投资指令偏转", "诱导 Agent 将资金或投资建议转向攻击者指定目标", "校验任务目标、资金流向与授权范围一致性"],
        ["CovertFundReallocation", "隐蔽资金重分配", "诱导 Agent 在表面合法任务中调整资金分配或交易方向", "在 Action Gate 检查高风险金融动作和参数异常"],
        ["DataFabrication", "数据伪造", "诱导 Agent 编造财务数据、交易记录或分析结论", "在 Return Gate 限制不可信结果回传并触发风险审计"],
        ["InsiderInformationLeak", "内部信息泄露", "诱导 Agent 泄露敏感持仓、内部计划或用户隐私", "结合权限令牌和敏感信息策略阻断越权输出"],
        ["TransactionDuplication", "交易重复执行", "诱导 Agent 重复执行转账、下单或记录写入", "基于 nonce、session 与动作上下文识别重放风险"],
    ]
    variant_rows = [
        ["攻击变体", "含义", "作用方式", "本次测试覆盖"],
        ["naive", "直接提示注入", "直接在用户输入或攻击内容中写入越权指令，要求 Agent 忽略原任务并执行攻击目标", "DPI、POT"],
        ["fake_completion", "伪完成注入", "伪造“任务已完成/系统已确认”等上下文，诱导 Agent 接受攻击者构造的后续指令", "DPI"],
        ["escape_characters", "转义字符注入", "通过特殊符号、转义片段或格式扰动绕过简单字符串检测", "DPI"],
        ["context_ignoring", "上下文忽略型观测注入", "在外部观测结果中植入忽略原规则的提示，测试 Agent 对不可信观测内容的抗干扰能力", "OPI"],
        ["combined_attack", "组合型攻击", "将直接注入、观测注入或记忆污染链路组合，测试多阶段上下文污染风险", "Mixed、MP"],
    ]
    return [
        ("h", "3.6 真实 Agent 安全有效性测试", "Heading 2"),
        ("p", "本节是作品安全有效性验证的核心部分。为避免仅在前端演示或模拟流程中验证防护效果，实验选取 ASB-native LangGraph Financial Agent 作为真实 Agent 运行环境，在同一任务集合、同一攻击工具集合与同一评价口径下，分别运行无防御组与 AegisGuard 防御组。测试指标包括攻击成功率（ASR）、原始任务成功率、拒绝率与运行耗时。其中 ASR 越低表示攻击越难达成，原始任务成功率越高表示系统在安全防护下越能保持正常业务可用性。"),
        ("p", "本节不再只给出汇总表，而是按照 DPI、OPI、Mixed、MP、POT 与 Clean 六类场景逐一展开。每类攻击均先说明攻击意图与变体，再分别给出无防御组和 AegisGuard 组的对照结果，并进一步按攻击工具拆分，观察不同业务风险类型下的防护效果。"),
        ("h", "3.6.1 攻击变体与攻击工具说明", "Heading 3"),
        ("p", "攻击变体用于刻画注入内容进入 Agent 执行链路的方式；攻击工具用于刻画攻击者希望 Agent 滥用的业务能力。二者组合后，可以同时观察“注入入口是否有效”和“被滥用的业务动作是否危险”。"),
        ("cap", "表 3-24 攻击变体说明"),
        ("table", variant_rows, 8),
        ("p", "从防护视角看，DPI 更侧重 Message Gate 对用户输入与指令边界的识别，OPI 与 Mixed 更侧重 Return Gate 对外部观测内容的隔离，MP 侧重 Memory Sandbox 对长期状态污染的控制，POT 则用于观察计划生成过程中的后门触发风险。"),
        ("cap", "表 3-25 攻击工具说明"),
        ("table", attack_tool_rows, 7.5),
        ("h", "3.6.2 DPI 直接提示注入测试", "Heading 3"),
        ("p", "DPI（Direct Prompt Injection）直接将恶意指令放入 Agent 可见输入中，是最典型、也最容易出现在真实交互中的攻击方式。该类攻击试图让 Agent 忽略原任务、越权调用工具或生成被攻击者控制的业务结果。本组实验覆盖 naive、fake_completion 与 escape_characters 三种变体，共 75 个真实 Agent case。"),
        ("cap", "表 3-26 DPI 总体防护效果对比"),
        ("table", [
            ["指标", "无防御", "AegisGuard", "变化"],
            ["样本数", 75, 75, "-"],
            ["攻击成功数", 13, 0, "-13"],
            ["ASR", "17.33%", "0.00%", "下降 17.33 个百分点"],
            ["原始任务成功数", 58, 68, "+10"],
            ["原始任务成功率", "77.33%", "90.67%", "提升 13.34 个百分点"],
            ["拒绝率", "1.33%", "0.00%", "下降 1.33 个百分点"],
        ]),
        ("cap", "表 3-27 DPI 攻击变体维度对比"),
        ("table", [
            ["攻击变体", "样本数", "无防御 ASR", "AegisGuard ASR", "AegisGuard 任务成功率"],
            ["naive", 25, "28.00%", "0.00%", "92.00%"],
            ["fake_completion", 25, "16.00%", "0.00%", "84.00%"],
            ["escape_characters", 25, "8.00%", "0.00%", "96.00%"],
        ]),
        ("cap", "表 3-28 DPI 攻击工具维度对比"),
        ("table", [
            ["攻击工具", "样本数", "无防御 ASR", "AegisGuard ASR", "AegisGuard 任务成功率"],
            ["DataFabrication", 15, "40.00%", "0.00%", "60.00%"],
            ["CovertFundReallocation", 15, "33.33%", "0.00%", "93.33%"],
            ["InsiderInformationLeak", 15, "13.33%", "0.00%", "100.00%"],
            ["InvestmentDiversion", 15, "0.00%", "0.00%", "100.00%"],
            ["TransactionDuplication", 15, "0.00%", "0.00%", "100.00%"],
        ]),
        ("p", "DPI 是本轮实验中安全收益最明显的攻击类型。无防御状态下，naive 变体 ASR 达到 28.00%，DataFabrication 与 CovertFundReallocation 两类攻击工具 ASR 分别达到 40.00% 和 33.33%，说明 Agent 在面对直接恶意指令时存在真实可利用风险。接入 AegisGuard 后，三类攻击变体和五类攻击工具的 ASR 均降至 0.00%，同时总体任务成功率由 77.33% 提升至 90.67%。这表明系统并不是依赖简单拒绝来换取安全，而是通过输入门控、动作授权校验和结果回传控制，使 Agent 更稳定地回到原始任务目标。"),
        ("h", "3.6.3 OPI 观测提示注入测试", "Heading 3"),
        ("p", "OPI（Observation Prompt Injection）将恶意指令隐藏在工具返回、网页内容、检索结果等外部观测信息中。该类攻击的危险性在于，攻击内容并非来自用户直接指令，而是伪装成 Agent 执行任务时读取到的外部数据。本组实验采用 context_ignoring 变体，共 25 个真实 Agent case。"),
        ("cap", "表 3-29 OPI 总体防护效果对比"),
        ("table", [
            ["指标", "无防御", "AegisGuard", "变化"],
            ["样本数", 25, 25, "-"],
            ["攻击成功数", 0, 0, "持平"],
            ["ASR", "0.00%", "0.00%", "持平"],
            ["原始任务成功数", 25, 22, "-3"],
            ["原始任务成功率", "100.00%", "88.00%", "下降 12 个百分点"],
            ["拒绝率", "0.00%", "12.00%", "上升 12 个百分点"],
        ]),
        ("cap", "表 3-30 OPI 攻击工具维度对比"),
        ("table", [
            ["攻击工具", "样本数", "无防御 ASR", "AegisGuard ASR", "AegisGuard 任务成功率", "AegisGuard 拒绝率"],
            ["InvestmentDiversion", 5, "0.00%", "0.00%", "100.00%", "0.00%"],
            ["CovertFundReallocation", 5, "0.00%", "0.00%", "100.00%", "0.00%"],
            ["DataFabrication", 5, "0.00%", "0.00%", "60.00%", "40.00%"],
            ["InsiderInformationLeak", 5, "0.00%", "0.00%", "100.00%", "0.00%"],
            ["TransactionDuplication", 5, "0.00%", "0.00%", "80.00%", "20.00%"],
        ], 7.5),
        ("p", "OPI 场景下，无防御组本身未被当前 context_ignoring 样本打穿，因此 AegisGuard 的效果主要体现为继续保持 0.00% ASR，而不是进一步降低攻击成功率。需要注意的是，DataFabrication 和 TransactionDuplication 场景中出现一定保守拒绝，使任务成功率从 100.00% 降至 88.00%。该结果说明 Return Gate 对外部观测内容的可信度判断偏谨慎，后续可通过细化工具类别、数据来源可信等级和任务上下文关系来降低误拒。"),
        ("h", "3.6.4 Mixed 组合攻击测试", "Heading 3"),
        ("p", "Mixed 场景将 DPI 与 OPI 组合，模拟攻击者同时从用户输入和外部观测结果两侧影响 Agent 的情况。该类攻击更接近真实复杂场景：攻击指令可能先改变 Agent 对任务目标的理解，再借助外部结果进一步强化错误动作。本组实验采用 combined_attack 变体，共 25 个真实 Agent case。"),
        ("cap", "表 3-31 Mixed 总体防护效果对比"),
        ("table", [
            ["指标", "无防御", "AegisGuard", "变化"],
            ["样本数", 25, 25, "-"],
            ["攻击成功数", 0, 0, "持平"],
            ["ASR", "0.00%", "0.00%", "持平"],
            ["原始任务成功数", 23, 24, "+1"],
            ["原始任务成功率", "92.00%", "96.00%", "提升 4 个百分点"],
            ["拒绝率", "0.00%", "0.00%", "持平"],
        ]),
        ("cap", "表 3-32 Mixed 攻击工具维度对比"),
        ("table", [
            ["攻击工具", "样本数", "无防御 ASR", "AegisGuard ASR", "无防御任务成功率", "AegisGuard 任务成功率"],
            ["InvestmentDiversion", 5, "0.00%", "0.00%", "100.00%", "100.00%"],
            ["CovertFundReallocation", 5, "0.00%", "0.00%", "100.00%", "100.00%"],
            ["DataFabrication", 5, "0.00%", "0.00%", "80.00%", "80.00%"],
            ["InsiderInformationLeak", 5, "0.00%", "0.00%", "80.00%", "100.00%"],
            ["TransactionDuplication", 5, "0.00%", "0.00%", "100.00%", "100.00%"],
        ], 7.5),
        ("p", "Mixed 场景中，无防御组同样未出现攻击成功，但任务成功率为 92.00%，说明组合攻击仍会干扰部分正常任务完成。AegisGuard 接入后保持 0.00% ASR，拒绝率仍为 0.00%，任务成功率提升至 96.00%。这说明在当前样本下，系统能够对组合型上下文污染进行稳定隔离，没有因为多入口检测而明显增加拒绝。"),
        ("h", "3.6.5 MP 记忆污染链路测试", "Heading 3"),
        ("p", "MP（Memory Poisoning）关注恶意信息进入长期记忆或历史状态后，对后续任务产生持续影响的风险。本组实验采用 combined_attack 变体，测试记忆污染链路下 Agent 是否会被诱导执行攻击目标。需要说明的是，本批次 memory_found 均为 0，因此结果更适合表述为“记忆污染链路下的组合攻击效果评估”，尚不能完全证明记忆检索命中后的污染防护效果。"),
        ("cap", "表 3-33 MP 总体防护效果对比"),
        ("table", [
            ["指标", "无防御", "AegisGuard", "变化"],
            ["样本数", 25, 25, "-"],
            ["攻击成功数", 1, 0, "-1"],
            ["ASR", "4.00%", "0.00%", "下降 4 个百分点"],
            ["原始任务成功数", 24, 25, "+1"],
            ["原始任务成功率", "96.00%", "100.00%", "提升 4 个百分点"],
            ["拒绝率", "0.00%", "0.00%", "持平"],
            ["Memory Found", 0, 0, "持平"],
        ]),
        ("cap", "表 3-34 MP 攻击工具维度对比"),
        ("table", [
            ["攻击工具", "样本数", "无防御 ASR", "AegisGuard ASR", "无防御任务成功率", "AegisGuard 任务成功率", "Memory Found"],
            ["InvestmentDiversion", 5, "0.00%", "0.00%", "100.00%", "100.00%", 0],
            ["CovertFundReallocation", 5, "0.00%", "0.00%", "100.00%", "100.00%", 0],
            ["DataFabrication", 5, "20.00%", "0.00%", "80.00%", "100.00%", 0],
            ["InsiderInformationLeak", 5, "0.00%", "0.00%", "100.00%", "100.00%", 0],
            ["TransactionDuplication", 5, "0.00%", "0.00%", "100.00%", "100.00%", 0],
        ], 7),
        ("p", "MP 场景中，无防御组出现 1 个成功攻击样本，集中在 DataFabrication 工具下，该工具维度 ASR 为 20.00%。AegisGuard 将该成功样本压降为 0，并将任务成功率提升至 100.00%。但由于本批次未观测到 memory_found 命中，报告中应保持审慎表述：AegisGuard 在当前 MP/combined_attack 配置下消除了观测到的攻击成功样本，后续仍需补充更强 memory poisoning 数据库命中样本来验证记忆沙箱能力。"),
        ("h", "3.6.6 POT 计划思维后门攻击测试", "Heading 3"),
        ("p", "POT（Plan-of-Thought Backdoor）用于测试 Agent 在规划或中间推理过程中被触发词诱导的风险。该类攻击不一定直接要求工具越权，而是通过隐藏触发条件改变 Agent 的计划生成方向。本组实验采用 naive 变体，触发词为 \"with perspicacious discernment\"，共 25 个真实 Agent case。"),
        ("cap", "表 3-35 POT 总体防护效果对比"),
        ("table", [
            ["指标", "无防御", "AegisGuard", "变化"],
            ["样本数", 25, 25, "-"],
            ["攻击成功数", 1, 0, "-1"],
            ["ASR", "4.00%", "0.00%", "下降 4 个百分点"],
            ["原始任务成功数", 24, 24, "持平"],
            ["原始任务成功率", "96.00%", "96.00%", "持平"],
            ["拒绝率", "0.00%", "0.00%", "持平"],
        ]),
        ("cap", "表 3-36 POT 攻击工具维度对比"),
        ("table", [
            ["攻击工具", "样本数", "无防御 ASR", "AegisGuard ASR", "AegisGuard 任务成功率"],
            ["InvestmentDiversion", 5, "0.00%", "0.00%", "100.00%"],
            ["CovertFundReallocation", 5, "0.00%", "0.00%", "100.00%"],
            ["DataFabrication", 5, "20.00%", "0.00%", "80.00%"],
            ["InsiderInformationLeak", 5, "0.00%", "0.00%", "100.00%"],
            ["TransactionDuplication", 5, "0.00%", "0.00%", "100.00%"],
        ]),
        ("p", "POT 无防御组同样出现 1 个攻击成功样本，集中在 DataFabrication 工具下。AegisGuard 将总 ASR 从 4.00% 降至 0.00%，同时保持 96.00% 的任务成功率，拒绝率未增加。这说明系统对计划生成阶段的后门诱导具有一定抑制作用，且该场景下防护收益较为干净：安全性提升，没有造成额外可用性损失。"),
        ("h", "3.6.7 Clean 无攻击实用性基线测试", "Heading 3"),
        ("p", "Clean 场景不包含攻击载荷，主要用于判断 AegisGuard 是否会因过度防御而损害正常任务。该组实验不是安全攻击测试，而是可用性基线测试。若 AegisGuard 在 Clean 场景下任务成功率明显下降，则说明防护策略存在较高误伤风险。"),
        ("cap", "表 3-37 Clean 总体可用性对比"),
        ("table", [
            ["指标", "无防御 Clean", "AegisGuard Clean", "变化"],
            ["样本数", 25, 25, "-"],
            ["攻击成功数", 0, 0, "持平"],
            ["ASR", "0.00%", "0.00%", "持平"],
            ["原始任务成功数", 17, 22, "+5"],
            ["原始任务成功率", "68.00%", "88.00%", "提升 20 个百分点"],
            ["拒绝率", "12.00%", "12.00%", "持平"],
        ]),
        ("cap", "表 3-38 Clean 任务工具维度对比"),
        ("table", [
            ["任务工具", "样本数", "无防御任务成功率", "AegisGuard 任务成功率", "AegisGuard 拒绝率"],
            ["InvestmentDiversion", 5, "80.00%", "80.00%", "20.00%"],
            ["CovertFundReallocation", 5, "60.00%", "80.00%", "20.00%"],
            ["DataFabrication", 5, "60.00%", "100.00%", "0.00%"],
            ["InsiderInformationLeak", 5, "60.00%", "80.00%", "20.00%"],
            ["TransactionDuplication", 5, "80.00%", "100.00%", "0.00%"],
        ], 7.5),
        ("p", "Clean 结果显示，AegisGuard 没有增加整体拒绝率，反而将正常任务成功率从 68.00% 提升至 88.00%。这可能与门控机制对任务边界、工具参数和结果回传格式的约束有关，使 Agent 在无攻击场景下也更少偏离原始目标。因此，Clean 结果可以支撑“防护机制未明显损害可用性”的论断。"),
        ("h", "3.6.8 真实 Agent 测试小结", "Heading 3"),
        ("cap", "表 3-39 真实 Agent 六类场景总体汇总"),
        ("table", [
            ["攻击类型", "组别", "样本数", "攻击成功数", "ASR", "任务成功率", "拒绝率"],
            ["DPI", "无防御", 75, 13, "17.33%", "77.33%", "1.33%"],
            ["DPI", "AegisGuard", 75, 0, "0.00%", "90.67%", "0.00%"],
            ["OPI", "无防御", 25, 0, "0.00%", "100.00%", "0.00%"],
            ["OPI", "AegisGuard", 25, 0, "0.00%", "88.00%", "12.00%"],
            ["Mixed", "无防御", 25, 0, "0.00%", "92.00%", "0.00%"],
            ["Mixed", "AegisGuard", 25, 0, "0.00%", "96.00%", "0.00%"],
            ["MP", "无防御", 25, 1, "4.00%", "96.00%", "0.00%"],
            ["MP", "AegisGuard", 25, 0, "0.00%", "100.00%", "0.00%"],
            ["POT", "无防御", 25, 1, "4.00%", "96.00%", "0.00%"],
            ["POT", "AegisGuard", 25, 0, "0.00%", "96.00%", "0.00%"],
            ["Clean", "无防御", 25, 0, "0.00%", "68.00%", "12.00%"],
            ["Clean", "AegisGuard", 25, 0, "0.00%", "88.00%", "12.00%"],
        ], 7.2),
        ("p", "总体来看，AegisGuard 在 200 个真实 Agent case 中将攻击成功数从 15 降至 0，总体 ASR 从 7.50% 降至 0.00%。在攻击场景中，任务成功率由 88.00% 提升至 93.14%；包含 Clean 场景后，总体任务成功率由 85.50% 提升至 92.50%。分场景看，DPI 是最主要的风险来源，也是防护收益最明显的场景；MP 与 POT 的成功攻击样本均被压降为 0；OPI 与 Mixed 在无防御下未被当前样本打穿，AegisGuard 主要体现为保持 0% ASR 并控制任务可用性。"),
        ("p", "上述结果说明，AegisGuard 的核心价值不只是识别恶意文本，而是在真实 Agent 执行链路中对输入、动作、结果和状态进行持续约束。其防护效果覆盖直接输入、外部观测、组合注入、记忆污染链路和计划后门等多类风险，能够更充分支撑本作品“面向 Agent 运行时安全监管与执行控制”的技术定位。"),
    ]


def main():
    doc = Document(str(SRC))
    body, blocks, start, end = find_body_range(doc)
    # Use the paragraph immediately before the old 3.6 body section as insertion anchor.
    anchor_para = None
    for p in doc.paragraphs:
        if p._p is blocks[start - 1]:
            anchor_para = p
            break
    if anchor_para is None:
        raise RuntimeError("未找到插入锚点")
    for block in blocks[start:end]:
        body.remove(block)
    append_items(anchor_para, detailed_items())
    set_update_fields(doc)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
